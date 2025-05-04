import tkinter as tk
from tkinter import messagebox
from docx import Document
import os
import re

TEMPLATE_PATH = r"C:\Users\z0052raf\Documents\Prototype Generator Contract Ver. 2\Draft Contract Template_TOP PLUS_CT.docx"

def sanitize_filename(name):
    return re.sub(r'[<>:"/\\|?*]', '', name)

def replace_placeholder_in_paragraph(paragraph, placeholder, replacement):
    full_text = ''.join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        new_text = full_text.replace(placeholder, replacement)
        for run in paragraph.runs:
            run.text = ''
        paragraph.runs[0].text = new_text

def replace_placeholder_in_table(table, placeholder, replacement):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                replace_placeholder_in_paragraph(para, placeholder, replacement)

def generate_contract():
    if not os.path.exists(TEMPLATE_PATH):
        messagebox.showerror("Error", f"Template not found:\n{TEMPLATE_PATH}")
        return

    doc = Document(TEMPLATE_PATH)

    ref_no        = entry_refno.get().strip()
    client        = entry_client_name.get().strip()
    addr1         = entry_address1.get().strip()
    addr2         = entry_address2.get().strip()
    initial       = entry_initial.get().strip()
    location      = entry_location.get().strip()
    equipment_raw = entry_equipment.get("1.0", "end-1c").strip()
    equipment     = "\n".join([ln for ln in equipment_raw.splitlines() if ln.strip()])
    duration      = entry_contract_duration.get().strip()
    start         = entry_start_contract.get().strip()
    end           = entry_end_contract.get().strip()
    pm            = entry_preventive.get().strip()
    price         = entry_price.get().strip()
    pricewords    = entry_pricewords.get().strip()
    peryear       = entry_per_year.get().strip()
    pricexyear    = entry_price_x_year.get().strip()
    pricewordsx   = entry_price_words_x.get().strip()
    periodeng     = entry_period_eng.get().strip()
    periodind     = entry_period_ind.get().strip()
    contractdate  = entry_contract_date.get().strip()
    sign1         = entry_signatory1.get().strip()
    pos1          = entry_pos1.get().strip()
    sign2         = entry_signatory2.get().strip()
    pos2          = entry_pos2.get().strip()
    sign3         = entry_signatory3.get().strip()
    pos3          = entry_pos3.get().strip()
    offereng      = entry_offereng.get().strip()
    offerind      = entry_offerind.get().strip()
    confeng       = entry_confeng.get().strip()
    confind       = entry_confind.get().strip()
    eqlist        = entry_eqlist.get().strip()
    sn            = entry_sn.get().strip()
    eqno          = entry_eqno.get().strip()

    if not ref_no or not client:
        messagebox.showerror("Error", "Ref No dan Client Name wajib diisi!")
        return

    placeholders = {
        "{{REFNO}}": ref_no,
        "{{CLIENT}}": client,
        "{{ADDRESS1}}": addr1,
        "{{ADDRESS2}}": addr2,
        "{{INITIAL}}": initial,
        "{{LOCATION}}": location,
        "{{EQUIPMENT}}": equipment,
        "{{DURATION}}": duration,
        "{{START}}": start,
        "{{END}}": end,
        "{{MAINTANANCE}}": pm,
        "{{PRICE}}": price,
        "{{PRICEWORDS}}": pricewords,
        "{{PERYEAR}}": peryear,
        "{{PRICEPERXYEAR}}": pricexyear,
        "{{PRICEWORDSX}}": pricewordsx,
        "{{PERIODENG}}": periodeng,
        "{{PERIODIND}}": periodind,
        "{{CONTRACTDATE}}": contractdate,
        "{{SIGNATORY1}}": sign1,
        "{{POS1}}": pos1,
        "{{SIGNATORY2}}": sign2,
        "{{POS2}}": pos2,
        "{{SIGNATORY3}}": sign3,
        "{{POS3}}": pos3,
        "{{OFFERENG}}": offereng,
        "{{OFFERIND}}": offerind,
        "{{CONFENG}}": confeng,
        "{{CONFIND}}": confind,
        "{{EQLIST}}": eqlist,
        "{{SN}}": sn,
        "{{EQNO}}": eqno
    }

    for para in doc.paragraphs:
        for ph, repl in placeholders.items():
            replace_placeholder_in_paragraph(para, ph, repl)
    for table in doc.tables:
        for ph, repl in placeholders.items():
            replace_placeholder_in_table(table, ph, repl)

    save_folder = os.path.dirname(TEMPLATE_PATH)
    fn = f"{sanitize_filename(ref_no)}_{sanitize_filename(client)}_Contract.docx"
    out_path = os.path.join(save_folder, fn)
    doc.save(out_path)
    messagebox.showinfo("Success", f"Contract generated!\nDisimpan di:\n{out_path}")

# GUI Setup
root = tk.Tk()
root.title("Contract Generator")
root.geometry("600x700")

# Add title label
container = tk.Frame(root)
container.pack(fill="both", expand=True)

canvas = tk.Canvas(container)
scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
canvas.configure(yscrollcommand=scrollbar.set)

scrollbar.pack(side="right", fill="y")
canvas.pack(side="left", fill="both", expand=True)

def on_mousewheel(event):
    if event.num == 4:
        canvas.yview_scroll(-1, "units")
    elif event.num == 5:
        canvas.yview_scroll(1, "units")

canvas.bind_all("<Button-4>", on_mousewheel)
canvas.bind_all("<Button-5>", on_mousewheel)


form = tk.Frame(canvas)

header_label = tk.Label(form, text="📄 Contract Generator Project", font=("Helvetica", 16, "bold"), fg="#003366")
header_label.grid(row=0, column=0, columnspan=2, pady=(10, 20))

canvas.create_window((0,0), window=form, anchor="nw")
form.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

fields = [
    ("== General Info ==", None, None),
    ("Ref No", tk.Entry, "entry_refno"),
    ("Client Name", tk.Entry, "entry_client_name"),
    ("Address Line 1", tk.Entry, "entry_address1"),
    ("Address Line 2", tk.Entry, "entry_address2"),
    ("Client Initial", tk.Entry, "entry_initial"),
    ("Location", tk.Entry, "entry_location"),
    ("Equipment", tk.Text, "entry_equipment", 5),

    ("== Contract Period ==", None, None),
    ("Contract Duration", tk.Entry, "entry_contract_duration"),
    ("Start Contract", tk.Entry, "entry_start_contract"),
    ("End Contract", tk.Entry, "entry_end_contract"),
    ("Preventive Maintenance", tk.Entry, "entry_preventive"),

    ("== Pricing ==", None, None),
    ("Price", tk.Entry, "entry_price"),
    ("Price in Words", tk.Entry, "entry_pricewords"),
    ("Per Year", tk.Entry, "entry_per_year"),
    ("Price per X Year", tk.Entry, "entry_price_x_year"),
    ("Price Words per X Year", tk.Entry, "entry_price_words_x"),

    ("== Period Language ==", None, None),
    ("Period ENG", tk.Entry, "entry_period_eng"),
    ("Period IND", tk.Entry, "entry_period_ind"),

    ("== Dates & References ==", None, None),
    ("Contract Date", tk.Entry, "entry_contract_date"),
    ("Offer Date ENG", tk.Entry, "entry_offereng"),
    ("Offer Date IND", tk.Entry, "entry_offerind"),
    ("Confirmation Date ENG", tk.Entry, "entry_confeng"),
    ("Confirmation Date IND", tk.Entry, "entry_confind"),
    ("Equipment List", tk.Entry, "entry_eqlist"),
    ("Serial Number", tk.Entry, "entry_sn"),
    ("Equipment No.", tk.Entry, "entry_eqno"),

    ("== Signatories ==", None, None),
    ("Signatory 1", tk.Entry, "entry_signatory1"),
    ("Position 1", tk.Entry, "entry_pos1"),
    ("Signatory 2", tk.Entry, "entry_signatory2"),
    ("Position 2", tk.Entry, "entry_pos2"),
    ("Signatory 3", tk.Entry, "entry_signatory3"),
    ("Position 3", tk.Entry, "entry_pos3"),
]

row_idx = 1  # start after the title header
for field in fields:
    label_text, widget_type, var_name, *opt = field

    if widget_type is None:
        # It's a section header
        section_label = tk.Label(form, text=label_text.replace("==", "").strip(), font=("Helvetica", 12, "bold"), fg="#004488")
        section_label.grid(row=row_idx, column=0, columnspan=2, sticky="w", pady=(15, 5), padx=5)
        row_idx += 1
        continue

    tk.Label(form, text=label_text).grid(row=row_idx, column=0, sticky="ne" if widget_type == tk.Text else "e", padx=5, pady=5)
    
    if widget_type == tk.Text:
        height = opt[0]
        w = tk.Text(form, width=50, height=height)
    else:
        w = widget_type(form, width=50)
    
    w.grid(row=row_idx, column=1, padx=5, pady=5)
    globals()[var_name] = w
    row_idx += 1

btn = tk.Button(form, text="Generate Contract", command=generate_contract, bg="lightblue", height=2)
btn.grid(row=row_idx, column=0, columnspan=2, pady=20)

root.mainloop()
